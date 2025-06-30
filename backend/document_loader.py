"""
Document Loader Module
Handles loading and parsing different document types (PDF, DOCX, TXT)
"""

import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from typing import List, Optional
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

class DocumentLoader:
    """Handles loading and processing different document types"""
    
    def __init__(self):
        """Initialize document loader with text splitter"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_document(self, file_path: str) -> Optional[List[Document]]:
        """Load document based on file extension"""
        try:
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._load_pdf(file_path)
            elif file_extension == 'docx':
                return self._load_docx(file_path)
            elif file_extension == 'txt':
                return self._load_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            print(f"Error loading document {file_path}: {str(e)}")
            return None
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load and parse PDF document"""
        try:
            doc = fitz.open(file_path)
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(text_chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        'source': os.path.basename(file_path),
                        'page': i + 1,
                        'file_type': 'pdf'
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            print(f"Error loading PDF {file_path}: {str(e)}")
            return []
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Load and parse DOCX document"""
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(text_chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        'source': os.path.basename(file_path),
                        'section': i + 1,
                        'file_type': 'docx'
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {str(e)}")
            return []
    
    def _load_txt(self, file_path: str) -> List[Document]:
        """Load and parse TXT document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(text_chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        'source': os.path.basename(file_path),
                        'section': i + 1,
                        'file_type': 'txt'
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            print(f"Error loading TXT {file_path}: {str(e)}")
            return []
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about a document"""
        try:
            file_size = os.path.getsize(file_path)
            file_extension = file_path.lower().split('.')[-1]
            
            info = {
                'filename': os.path.basename(file_path),
                'file_size': file_size,
                'file_type': file_extension,
                'file_path': file_path
            }
            
            # Get additional info based on file type
            if file_extension == 'pdf':
                doc = fitz.open(file_path)
                info['pages'] = len(doc)
                doc.close()
            elif file_extension == 'docx':
                doc = DocxDocument(file_path)
                info['paragraphs'] = len(doc.paragraphs)
            elif file_extension == 'txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    lines = file.readlines()
                    info['lines'] = len(lines)
            
            return info
            
        except Exception as e:
            print(f"Error getting document info for {file_path}: {str(e)}")
            return {} 